"""
Creates a sample Word document (.docx) with various formatting structures
for testing the content migration pipeline.
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def add_hyperlink(paragraph, url, text):
    """Add a hyperlink to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")

    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    rPr.append(color)

    u = OxmlElement("w:u")
    u.set(qn("w:val"), "single")
    rPr.append(u)

    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

    return hyperlink


def create_sample_document(filename="sample_document.docx"):
    """Create a comprehensive sample Word document."""
    doc = Document()

    # ── Heading 1 ──
    doc.add_heading("Content Migration Guide", level=1)

    # ── Introductory paragraph ──
    doc.add_paragraph(
        "This document demonstrates various formatting structures commonly found "
        "in Microsoft Word documents. The content migration tool should be able to "
        "parse and convert all of these elements into clean HTML."
    )

    # ── Heading 2 ──
    doc.add_heading("Getting Started", level=2)

    doc.add_paragraph(
        "Before beginning the migration process, ensure you have the necessary "
        "API credentials and access to the target platform."
    )

    # ── Heading 3 ──
    doc.add_heading("Prerequisites", level=3)

    # ── Bullet list ──
    doc.add_paragraph("Python 3.8 or higher", style="List Bullet")
    doc.add_paragraph("A valid Document360 API token", style="List Bullet")
    doc.add_paragraph("The python-docx library installed", style="List Bullet")
    doc.add_paragraph("Access to the target knowledge base", style="List Bullet")

    # ── Heading 2 ──
    doc.add_heading("Installation Steps", level=2)

    # ── Numbered list ──
    doc.add_paragraph("Clone the repository from GitHub.", style="List Number")
    doc.add_paragraph("Install dependencies using pip install -r requirements.txt.", style="List Number")
    doc.add_paragraph("Configure the .env file with your API credentials.", style="List Number")
    doc.add_paragraph("Run the application using python main.py.", style="List Number")

    # ── Another paragraph ──
    doc.add_paragraph(
        "Once the setup is complete, you can start migrating your Word documents "
        "to Document360 with a single command."
    )

    # ── Heading 2 ──
    doc.add_heading("Supported Features", level=2)

    doc.add_paragraph(
        "The migration tool supports the following Word document elements:"
    )

    # ── Table ──
    table = doc.add_table(rows=7, cols=3)
    table.style = "Table Grid"

    # Header row
    headers = ["Word Element", "HTML Output", "Status"]
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True

    # Data rows
    data = [
        ["Headings (H1-H6)", "<h1> to <h6>", "Supported"],
        ["Paragraphs", "<p>", "Supported"],
        ["Bullet Lists", "<ul><li>", "Supported"],
        ["Numbered Lists", "<ol><li>", "Supported"],
        ["Tables", "<table>", "Supported"],
        ["Hyperlinks", "<a href='...'>", "Supported"],
    ]

    for row_idx, row_data in enumerate(data, start=1):
        for col_idx, cell_text in enumerate(row_data):
            table.rows[row_idx].cells[col_idx].text = cell_text

    # ── Heading 2 ──
    doc.add_heading("Useful Resources", level=2)

    # ── Paragraph with hyperlinks ──
    p = doc.add_paragraph("For more information, visit the ")
    add_hyperlink(p, "https://apidocs.document360.io", "Document360 API Documentation")
    p.add_run(" or check out the ")
    add_hyperlink(p, "https://python-docx.readthedocs.io", "python-docx Documentation")
    p.add_run(".")

    p2 = doc.add_paragraph("You can also visit ")
    add_hyperlink(p2, "https://github.com", "GitHub")
    p2.add_run(" for the source code repository.")

    # ── Heading 2 ──
    doc.add_heading("Advanced Configuration", level=2)

    doc.add_heading("Environment Variables", level=3)

    doc.add_paragraph(
        "The application uses environment variables for configuration. "
        "These can be set in a .env file or directly in your system environment."
    )

    # ── Another bullet list ──
    doc.add_paragraph("DOCUMENT360_API_TOKEN – Your API authentication token", style="List Bullet")
    doc.add_paragraph("DOCUMENT360_BASE_URL – The API base URL for your region", style="List Bullet")
    doc.add_paragraph("DOCUMENT360_PROJECT_VERSION_ID – Target project version", style="List Bullet")
    doc.add_paragraph("DOCUMENT360_CATEGORY_ID – Category for new articles", style="List Bullet")

    doc.add_heading("Error Handling", level=3)

    doc.add_paragraph(
        "The application includes comprehensive error handling for common scenarios:"
    )

    doc.add_paragraph("File not found or invalid format", style="List Number")
    doc.add_paragraph("API authentication failures", style="List Number")
    doc.add_paragraph("Network connectivity issues", style="List Number")
    doc.add_paragraph("Rate limiting responses", style="List Number")

    # ── Conclusion ──
    doc.add_heading("Conclusion", level=2)

    doc.add_paragraph(
        "This content migration tool provides a streamlined way to move content "
        "from Word documents to Document360. By following the steps outlined above, "
        "you can automate the process of converting and uploading your documentation."
    )

    doc.save(filename)
    print(f"✅ Sample document created: {filename}")
    return filename


if __name__ == "__main__":
    create_sample_document()
