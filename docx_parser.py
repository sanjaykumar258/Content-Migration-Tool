"""
docx_parser.py
──────────────
Parses a .docx file and extracts structured content elements:
  - Headings (H1–H6)
  - Paragraphs
  - Bullet lists
  - Numbered lists
  - Tables (with header detection)
  - Hyperlinks (inline within any text element)
"""

from __future__ import annotations

import logging
import os
from dataclasses import dataclass, field
from enum import Enum, auto
from typing import List, Optional

from docx import Document
from docx.oxml.ns import qn
from docx.table import Table as DocxTable

logger = logging.getLogger(__name__)


# ─── Element types ───────────────────────────────────────────────────────────

class ElementType(Enum):
    HEADING = auto()
    PARAGRAPH = auto()
    BULLET_LIST = auto()
    NUMBERED_LIST = auto()
    TABLE = auto()
    IMAGE = auto()


# ─── Data classes representing extracted content ─────────────────────────────

@dataclass
class InlineContent:
    """Represents a segment of inline content (text, hyperlink, bold, italic)."""
    text: str
    bold: bool = False
    italic: bool = False
    hyperlink_url: Optional[str] = None
    is_image: bool = False
    image_id: Optional[str] = None


@dataclass
class ContentElement:
    """A single structural element extracted from the document."""
    element_type: ElementType
    level: int = 0                              # heading level (1-6)
    inline_contents: List[InlineContent] = field(default_factory=list)
    items: List[List[InlineContent]] = field(default_factory=list)  # list items
    table_data: List[List[str]] = field(default_factory=list)       # rows × cols
    has_header_row: bool = False
    image_data: Optional[bytes] = None                              # Binary image data
    image_ext: Optional[str] = None                                # .png, .jpg, etc.


# ─── Parser ──────────────────────────────────────────────────────────────────

class DocxParser:
    """Reads a .docx file and yields a flat list of `ContentElement` objects."""

    def __init__(self, filepath: str) -> None:
        if not os.path.isfile(filepath):
            raise FileNotFoundError(f"File not found: {filepath}")
        if not filepath.lower().endswith(".docx"):
            raise ValueError(f"Invalid file format (expected .docx): {filepath}")

        self.filepath = filepath
        self.document = Document(filepath)
        logger.info("Loaded document: %s", filepath)

    # ── public API ───────────────────────────────────────────────────────────

    def parse(self) -> List[ContentElement]:
        """Parse the document and return a list of structural content elements."""
        elements: List[ContentElement] = []
        current_list: Optional[ContentElement] = None
        current_list_type: Optional[ElementType] = None

        for block in self._iter_blocks():
            if isinstance(block, DocxTable):
                # Flush any open list
                if current_list:
                    elements.append(current_list)
                    current_list = None
                    current_list_type = None
                elements.append(self._parse_table(block))
                continue

            # It's a paragraph
            para = block
            style_name = (para.style.name or "").lower()

            # Determine element type
            etype = self._classify_paragraph(style_name)

            # If it is a list item, accumulate
            if etype in (ElementType.BULLET_LIST, ElementType.NUMBERED_LIST):
                if current_list_type != etype:
                    # Flush previous list if type changed
                    if current_list:
                        elements.append(current_list)
                    current_list = ContentElement(element_type=etype)
                    current_list_type = etype
                current_list.items.append(self._extract_inline(para))
                continue

            # Not a list item → flush any open list
            if current_list:
                elements.append(current_list)
                current_list = None
                current_list_type = None

            if etype == ElementType.HEADING:
                level = self._heading_level(style_name)
                elem = ContentElement(
                    element_type=ElementType.HEADING,
                    level=level,
                    inline_contents=self._extract_inline(para),
                )
                elements.append(elem)

            elif etype == ElementType.PARAGRAPH:
                text = para.text.strip()
                if not text:
                    continue  # skip blank paragraphs
                elem = ContentElement(
                    element_type=ElementType.PARAGRAPH,
                    inline_contents=self._extract_inline(para),
                )
                elements.append(elem)

        # Flush trailing list
        if current_list:
            elements.append(current_list)

        logger.info("Extracted %d content elements", len(elements))
        return elements

    # ── private helpers ──────────────────────────────────────────────────────

    def _iter_blocks(self):
        """Iterate document body children, yielding Paragraph or Table objects
        in document order."""
        body = self.document.element.body
        paragraphs = self.document.paragraphs
        tables = self.document.tables
        para_index = 0
        table_index = 0

        for child in body:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
            if tag == "p":
                if para_index < len(paragraphs):
                    yield paragraphs[para_index]
                    para_index += 1
            elif tag == "tbl":
                if table_index < len(tables):
                    yield tables[table_index]
                    table_index += 1

    @staticmethod
    def _classify_paragraph(style_name: str) -> ElementType:
        if "heading" in style_name:
            return ElementType.HEADING
        if "list bullet" in style_name:
            return ElementType.BULLET_LIST
        if "list number" in style_name:
            return ElementType.NUMBERED_LIST
        return ElementType.PARAGRAPH

    @staticmethod
    def _heading_level(style_name: str) -> int:
        """Extract heading level from style name like 'heading 2'."""
        for part in style_name.split():
            if part.isdigit():
                return min(int(part), 6)
        return 1  # default to H1

    def _extract_inline(self, paragraph) -> List[InlineContent]:
        """Extract inline content (runs, hyperlinks, and images) from a paragraph."""
        contents: List[InlineContent] = []

        for child in paragraph._p:
            tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag

            if tag == "hyperlink":
                url = self._get_hyperlink_url(child, paragraph)
                text = "".join(
                    node.text or ""
                    for node in child.iter()
                    if node.tag.endswith("}t") or node.tag == "w:t"
                )
                if text:
                    contents.append(InlineContent(text=text, hyperlink_url=url))

            elif tag == "r":
                # Regular run
                run_text = ""
                # Check for images in this run (modern drawings and legacy picts)
                for drawing in child.iter(qn("w:drawing")):
                    image_id = self._get_image_id(drawing)
                    if image_id:
                        contents.append(InlineContent(text="", is_image=True, image_id=image_id))

                for pict in child.iter(qn("w:pict")):
                    image_id = self._get_pict_id(pict)
                    if image_id:
                        contents.append(InlineContent(text="", is_image=True, image_id=image_id))

                for t_elem in child.iter():
                    t_tag = t_elem.tag.split("}")[-1] if "}" in t_elem.tag else t_elem.tag
                    if t_tag == "t" and t_elem.text:
                        run_text += t_elem.text

                if run_text:
                    bold = False
                    italic = False
                    rPr = child.find(qn("w:rPr"))
                    if rPr is not None:
                        bold = rPr.find(qn("w:b")) is not None
                        italic = rPr.find(qn("w:i")) is not None
                    contents.append(InlineContent(text=run_text, bold=bold, italic=italic))

        return contents

    def _get_image_id(self, drawing_elem) -> Optional[str]:
        """Extract the relationship ID of an image from a drawing element."""
        # This is a bit complex as MS Word nesting can vary
        for blip in drawing_elem.iter(qn("a:blip")):
            r_id = blip.get(qn("r:embed"))
            if r_id:
                return r_id
        return None

    def _get_pict_id(self, pict_elem) -> Optional[str]:
        """Extract the relationship ID of an image from a legacy pict element."""
        # Legacy images often use v:imagedata
        for imagedata in pict_elem.iter(qn("v:imagedata")):
            r_id = imagedata.get(qn("r:id"))
            if r_id:
                return r_id
        return None

    def get_image_payload(self, r_id: str) -> tuple[Optional[bytes], Optional[str]]:
        """Retrieve the binary image data and extension from a relationship ID."""
        try:
            image_part = self.document.part.related_parts[r_id]
            return image_part.blob, image_part.extension
        except Exception:
            return None, None

    def _get_hyperlink_url(self, hyperlink_elem, paragraph) -> Optional[str]:
        """Resolve a w:hyperlink element to its URL."""
        r_id = hyperlink_elem.get(qn("r:id"))
        if r_id:
            try:
                rel = paragraph.part.rels.get(r_id)
                if rel:
                    return rel._target
            except Exception:
                pass
        return None

    def _parse_table(self, table: DocxTable) -> ContentElement:
        """Parse a table into a ContentElement."""
        rows_data: List[List[str]] = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            rows_data.append(row_data)

        has_header = len(rows_data) > 1  # assume first row is header if >1 rows

        return ContentElement(
            element_type=ElementType.TABLE,
            table_data=rows_data,
            has_header_row=has_header,
        )
